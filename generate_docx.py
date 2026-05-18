# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 标题样式
for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.name = '黑体'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(14)
doc.styles['Heading 4'].font.size = Pt(12)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_bold_para(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('合家云社区物业管理平台')
run.font.size = Pt(28)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('需求分析报告')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_items = ['项目名称：合家云社区物业管理平台', '项目类型：数据库应用系统开发（B/S架构）',
              '开发环境：Spring Boot + Vue 3 + MySQL 8.0', '报告日期：2026年5月']
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)

doc.add_page_break()

# ========== 正文 ==========
doc.add_heading('1. 项目背景', level=1)

doc.add_heading('1.1 项目名称', level=2)
doc.add_paragraph('合家云社区物业管理平台')

doc.add_heading('1.2 项目类型', level=2)
doc.add_paragraph('数据库应用系统开发（B/S架构）')

doc.add_heading('1.3 开发目的', level=2)
doc.add_paragraph(
    '为社区物业管理部门设计的管理信息系统，实现对社区资产（小区、楼栋、单元、房屋）和业主信息的统一管理，'
    '支持多角色权限控制和数据统计分析，提高物业管理效率与信息化水平。')

doc.add_heading('1.4 开发环境', level=2)
add_table(
    ['类别', '技术选型'],
    [
        ['操作系统', 'Microsoft Windows 11'],
        ['数据库', 'MySQL 8.0'],
        ['后端框架', 'JDK 21 + Spring Boot 3.2.0 + MyBatis-Plus 3.5.5'],
        ['前端框架', 'Vue 3.4 + Vite 5 + Element Plus 2.x'],
        ['接口规范', 'REST API + JDBC'],
        ['构建工具', 'Maven (后端) / npm (前端)'],
        ['版本控制', 'Git + GitHub'],
    ]
)

# ========== 2. 功能需求 ==========
doc.add_heading('2. 功能需求', level=1)

doc.add_heading('2.1 角色划分', level=2)
doc.add_paragraph('系统采用 RBAC（基于角色的访问控制）模型，划分为两种角色：')
add_table(
    ['角色', '角色标识', '权限说明'],
    [
        ['管理员', 'admin', '拥有系统所有权限，可进行增删改查全部操作'],
        ['普通用户', 'user', '仅可查看数据，无任何写操作权限'],
    ]
)

add_bold_para('默认测试账号：')
doc.add_paragraph('管理员：admin / 123456', style='List Bullet')
doc.add_paragraph('普通用户：operator / 123456', style='List Bullet')

add_bold_para('权限控制机制：')
doc.add_paragraph('前端通过 localStorage 存储用户角色标识 (roleKey)', style='List Bullet')
doc.add_paragraph('后端通过请求头 X-User-Role 校验写操作权限', style='List Bullet')
doc.add_paragraph('前端根据角色动态显示/隐藏操作按钮', style='List Bullet')
doc.add_paragraph('路由守卫拦截未登录用户跳转至登录页', style='List Bullet')

# ---------- 2.2 功能模块 ----------
doc.add_heading('2.2 功能模块详细说明', level=2)

# --- 登录模块 ---
doc.add_heading('2.2.1 登录模块', level=3)
add_table(
    ['功能项', '功能说明'],
    [
        ['用户登录', '输入用户名和密码进行身份验证'],
        ['登录校验', '验证用户名、密码及账号状态（禁用账号无法登录）'],
        ['角色识别', '登录成功后返回用户ID、用户名、真实姓名、角色标识'],
        ['路由守卫', '未登录用户自动跳转登录页；已登录用户访问登录页自动跳转首页'],
        ['退出登录', '清除本地存储的用户信息，跳转至登录页'],
    ]
)
add_bold_para('登录接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/login', 'POST', '用户登录，返回用户信息及角色'],
        ['/api/user/info/{id}', 'GET', '获取用户详情（隐藏密码）'],
    ]
)

# --- 系统管理 ---
doc.add_heading('2.2.2 系统管理模块', level=3)

doc.add_heading('（1）用户管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['用户列表', '分页展示用户信息，支持按用户名、手机号模糊查询', '所有用户'],
        ['新增用户', '添加用户，设置用户名、密码、真实姓名、手机号、邮箱、状态', '管理员'],
        ['编辑用户', '修改用户基本信息（用户名不可修改）', '管理员'],
        ['删除用户', '删除用户及其角色关联关系', '管理员'],
        ['状态管理', '启用/禁用用户账号', '管理员'],
        ['角色分配', '为用户分配一个或多个角色', '管理员'],
        ['admin保护', 'admin账号仅能由admin自身修改/删除/分配角色', '系统约束'],
    ]
)
add_bold_para('用户管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/system/user/list', 'GET', '用户分页列表（支持username、phone查询）'],
        ['/api/system/user/{id}', 'GET', '根据ID查询用户'],
        ['/api/system/user', 'POST', '新增用户'],
        ['/api/system/user', 'PUT', '更新用户信息'],
        ['/api/system/user/{id}', 'DELETE', '删除用户（级联删除角色关联）'],
        ['/api/system/user/status/{id}', 'PUT', '更新用户状态'],
        ['/api/system/user/{id}/roles', 'GET', '查询用户角色列表'],
        ['/api/system/user/{id}/roles', 'PUT', '分配用户角色'],
    ]
)

doc.add_heading('（2）角色管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['角色列表', '分页展示角色信息，支持按角色名称模糊查询', '所有用户'],
        ['新增角色', '添加角色，设置角色名称、标识、描述、状态', '管理员'],
        ['编辑角色', '修改角色信息', '管理员'],
        ['删除角色', '删除角色', '管理员'],
        ['获取全部角色', '不分页获取所有角色（用于下拉选择）', '所有用户'],
    ]
)
add_bold_para('角色管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/system/role/list', 'GET', '角色分页列表（支持roleName查询）'],
        ['/api/system/role/all', 'GET', '获取全部角色列表'],
        ['/api/system/role/{id}', 'GET', '根据ID查询角色'],
        ['/api/system/role', 'POST', '新增角色'],
        ['/api/system/role', 'PUT', '更新角色'],
        ['/api/system/role/{id}', 'DELETE', '删除角色'],
    ]
)

# --- 社区资产管理 ---
doc.add_heading('2.2.3 社区资产管理模块', level=3)
doc.add_paragraph('社区资产采用四级层级结构：小区 → 楼栋 → 单元 → 房屋')

doc.add_heading('（1）小区管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['小区列表', '分页展示小区信息，支持按名称模糊查询', '所有用户'],
        ['新增小区', '添加小区，设置名称、地址、占地面积、状态', '管理员'],
        ['编辑小区', '修改小区信息', '管理员'],
        ['删除小区', '删除小区（需先删除下属楼栋）', '管理员'],
        ['资产统计', '统计各小区的楼栋数和房间数', '所有用户'],
    ]
)
add_bold_para('小区管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/community/list', 'GET', '小区分页列表（支持name查询）'],
        ['/api/community/all', 'GET', '获取全部小区列表'],
        ['/api/community/{id}', 'GET', '根据ID查询小区'],
        ['/api/community', 'POST', '新增小区'],
        ['/api/community', 'PUT', '更新小区'],
        ['/api/community/{id}', 'DELETE', '删除小区'],
        ['/api/community/statistics', 'GET', '小区资产统计（楼栋数、房间数）'],
    ]
)

doc.add_heading('（2）楼栋管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['楼栋列表', '分页展示楼栋信息，支持按小区过滤、按名称模糊查询，显示所属小区名称', '所有用户'],
        ['新增楼栋', '添加楼栋，选择所属小区，设置名称、楼层数、备注', '管理员'],
        ['编辑楼栋', '修改楼栋信息', '管理员'],
        ['删除楼栋', '删除楼栋（需先删除下属单元）', '管理员'],
        ['级联查询', '支持按小区ID过滤楼栋列表（用于下拉联动）', '所有用户'],
    ]
)
add_bold_para('楼栋管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/community/building/list', 'GET', '楼栋分页列表（支持communityId、name查询）'],
        ['/api/community/building/all', 'GET', '获取全部楼栋（支持communityId过滤）'],
        ['/api/community/building/{id}', 'GET', '根据ID查询楼栋'],
        ['/api/community/building', 'POST', '新增楼栋'],
        ['/api/community/building', 'PUT', '更新楼栋'],
        ['/api/community/building/{id}', 'DELETE', '删除楼栋'],
    ]
)

doc.add_heading('（3）单元管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['单元列表', '分页展示单元信息，支持按楼栋过滤、按名称模糊查询，显示所属小区、楼栋名称', '所有用户'],
        ['新增单元', '添加单元，选择所属楼栋，设置名称、楼层数', '管理员'],
        ['编辑单元', '修改单元信息', '管理员'],
        ['删除单元', '删除单元（需先删除下属房屋）', '管理员'],
        ['级联查询', '支持按楼栋ID过滤单元列表（用于下拉联动）', '所有用户'],
    ]
)
add_bold_para('单元管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/community/unit/list', 'GET', '单元分页列表（支持buildingId、name查询）'],
        ['/api/community/unit/all', 'GET', '获取全部单元（支持buildingId过滤）'],
        ['/api/community/unit/{id}', 'GET', '根据ID查询单元'],
        ['/api/community/unit', 'POST', '新增单元'],
        ['/api/community/unit', 'PUT', '更新单元'],
        ['/api/community/unit/{id}', 'DELETE', '删除单元'],
    ]
)

doc.add_heading('（4）房屋管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['房屋列表', '分页展示房屋信息，支持三级联动筛选（小区→楼栋→单元）、房间号模糊查询、状态筛选，显示完整层级路径', '所有用户'],
        ['新增房屋', '添加房屋，选择所属单元，设置房间号、楼层、面积、户型、状态', '管理员'],
        ['编辑房屋', '修改房屋信息', '管理员'],
        ['删除房屋', '删除房屋（已入住状态的房屋不可删除）', '管理员'],
        ['状态管理', '房屋状态分三种：空置(vacant)、已入住(occupied)、装修中(decorated)', '管理员'],
        ['状态统计', '统计各状态房屋数量及占比', '所有用户'],
    ]
)
add_bold_para('房屋管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/community/room/list', 'GET', '房屋分页列表（支持多条件筛选）'],
        ['/api/community/room/all', 'GET', '获取全部房屋（支持unitId过滤）'],
        ['/api/community/room/{id}', 'GET', '根据ID查询房屋'],
        ['/api/community/room', 'POST', '新增房屋'],
        ['/api/community/room', 'PUT', '更新房屋'],
        ['/api/community/room/{id}', 'DELETE', '删除房屋'],
        ['/api/community/room/statistics/status', 'GET', '房屋状态统计'],
    ]
)

# --- 业主管理 ---
doc.add_heading('2.2.4 业主管理模块', level=3)

doc.add_heading('（1）业主档案管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['业主列表', '分页展示业主信息，支持按姓名、手机号、身份证号模糊查询', '所有用户'],
        ['新增业主', '添加业主，设置姓名、手机号、身份证号、性别、车牌号、紧急联系人等', '管理员'],
        ['编辑业主', '修改业主信息', '管理员'],
        ['删除业主', '删除业主（级联删除人房绑定关系）', '管理员'],
        ['综合搜索', '支持按姓名、手机号、身份证号多字段模糊搜索', '所有用户'],
    ]
)
add_bold_para('业主管理接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/owner/list', 'GET', '业主分页列表（支持name、phone、idCard查询）'],
        ['/api/owner/all', 'GET', '获取全部业主列表'],
        ['/api/owner/{id}', 'GET', '根据ID查询业主'],
        ['/api/owner', 'POST', '新增业主'],
        ['/api/owner', 'PUT', '更新业主'],
        ['/api/owner/{id}', 'DELETE', '删除业主'],
        ['/api/owner/search', 'GET', '综合搜索（匹配姓名/手机号/身份证号）'],
    ]
)

doc.add_heading('（2）人房绑定管理', level=4)
add_table(
    ['功能项', '功能说明', '操作权限'],
    [
        ['绑定列表', '分页展示绑定关系，显示业主姓名、联系电话、完整房屋路径、关系类型、绑定时间', '所有用户'],
        ['新增绑定', '选择业主和房屋建立绑定关系，房屋选择采用四级级联选择器（小区→楼栋→单元→房屋）', '管理员'],
        ['解绑', '解除绑定关系（软删除，记录解绑时间）', '管理员'],
        ['关系类型', '支持三种关系：业主(owner)、家属(family)、租客(tenant)', '系统约束'],
        ['重复绑定检查', '同一业主与同一房屋不能重复绑定', '系统约束'],
        ['租客唯一约束', '同一房屋只能绑定一个租客', '系统约束'],
        ['状态联动', '绑定为业主关系时房屋状态自动变为"已入住"；解绑后若无其他绑定，房屋状态自动恢复为"空置"', '系统约束'],
        ['按业主查询', '支持根据业主ID查询其所有绑定房屋', '所有用户'],
    ]
)
add_bold_para('人房绑定接口：')
add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/owner/room/list', 'GET', '绑定分页列表（支持ownerId、roomId、relationType查询）'],
        ['/api/owner/room/owner/{ownerId}', 'GET', '查询指定业主的所有绑定房屋'],
        ['/api/owner/room', 'POST', '新增绑定（含重复检查、租客唯一检查、状态联动）'],
        ['/api/owner/room/{id}', 'DELETE', '解绑（软删除，含状态联动）'],
    ]
)

# --- 数据统计 ---
doc.add_heading('2.2.5 数据统计模块（首页仪表盘）', level=3)
add_table(
    ['功能项', '功能说明'],
    [
        ['概览数据', '显示小区总数、楼栋总数、房屋总数、业主总数'],
        ['小区资产统计表', '以表格形式展示各小区的名称、楼栋数、房间数'],
        ['房屋状态分布', '展示空置房、已入住、装修中的数量及百分比'],
    ]
)

# --- 查询需求 ---
doc.add_heading('2.3 数据查询需求', level=2)
doc.add_paragraph('各管理模块统一支持以下查询方式：')
add_table(
    ['查询类型', '说明', '实现方式'],
    [
        ['分页查询', '默认每页10条，支持切换10/20/50条每页', 'MyBatis-Plus分页插件'],
        ['模糊查询', '按名称等字段前缀/包含匹配', 'SQL LIKE语句'],
        ['精确查询', '按状态等字段精确匹配', 'SQL等值查询'],
        ['条件组合', '支持多条件同时筛选', 'LambdaQueryWrapper链式调用'],
        ['层级过滤', '楼栋按小区、单元按楼栋、房屋按单元/楼栋/小区过滤', '外键关联查询'],
        ['级联选择', '房屋选择采用四级级联选择器', '前端Cascader组件'],
    ]
)

# ========== 3. 数据需求 ==========
doc.add_heading('3. 数据需求', level=1)

doc.add_heading('3.1 数据库表清单', level=2)
doc.add_paragraph('系统共设计 9 张数据表，分为三个维度：')
add_table(
    ['维度', '序号', '表名', '说明', '主要字段'],
    [
        ['资产维度', '1', 'hjy_community', '小区表', 'id, name, address, area, status'],
        ['资产维度', '2', 'hjy_building', '楼栋表', 'id, community_id(FK), name, floors, remark'],
        ['资产维度', '3', 'hjy_unit', '单元表', 'id, building_id(FK), name, floor_num'],
        ['资产维度', '4', 'hjy_room', '房屋表', 'id, unit_id(FK), room_no, floor, area, room_type, status'],
        ['权限维度', '5', 'sys_user', '用户表', 'id, username, password, real_name, phone, email, status'],
        ['权限维度', '6', 'sys_role', '角色表', 'id, role_name, role_key, description, status'],
        ['权限维度', '7', 'sys_user_role', '用户角色关联表', 'id, user_id(FK), role_id(FK)'],
        ['业主维度', '8', 'hjy_owner', '业主档案表', 'id, name, phone, id_card, gender, birthday, car_plate'],
        ['业主维度', '9', 'hjy_owner_room', '人房关联表', 'id, owner_id(FK), room_id(FK), relation_type, bind_time'],
    ]
)

doc.add_heading('3.2 数据关系模型', level=2)
doc.add_paragraph('层级关系链：小区(1) ──(N)── 楼栋(1) ──(N)── 单元(1) ──(N)── 房屋(1) ──(N)── 人房绑定(N) ──(1)── 业主')
doc.add_paragraph('权限关系链：用户(N) ──(N)── 角色，通过用户角色关联表实现多对多关系')
doc.add_paragraph('人房绑定关系：支持业主(owner)、家属(family)、租客(tenant)三种关系类型')

doc.add_heading('3.3 数据完整性约束', level=2)
add_table(
    ['约束类型', '实现方式', '具体说明'],
    [
        ['主键约束', 'AUTO_INCREMENT', '每张表 id 为主键自增'],
        ['外键约束', 'ON DELETE RESTRICT / CASCADE', '资产层级表使用RESTRICT防止误删；关联表使用CASCADE级联删除'],
        ['唯一约束', 'UNIQUE KEY', '用户名唯一、同一业主-房屋-关系类型组合唯一'],
        ['非空约束', 'NOT NULL', '必填字段设置NOT NULL（如name、phone等）'],
        ['默认值', 'DEFAULT', '状态字段默认值（status默认1）、时间字段默认当前时间'],
        ['业务约束', '应用层校验', '同一房屋只能有一个租客；已入住房屋不可删除；admin账号受保护'],
    ]
)

doc.add_heading('3.4 外键约束详情', level=2)
add_table(
    ['外键名', '源表.字段', '目标表.字段', '删除策略'],
    [
        ['fk_building_community', 'hjy_building.community_id', 'hjy_community.id', 'RESTRICT'],
        ['fk_unit_building', 'hjy_unit.building_id', 'hjy_building.id', 'RESTRICT'],
        ['fk_room_unit', 'hjy_room.unit_id', 'hjy_unit.id', 'RESTRICT'],
        ['fk_user_role_user', 'sys_user_role.user_id', 'sys_user.id', 'CASCADE'],
        ['fk_user_role_role', 'sys_user_role.role_id', 'sys_role.id', 'CASCADE'],
        ['fk_or_owner', 'hjy_owner_room.owner_id', 'hjy_owner.id', 'CASCADE'],
        ['fk_or_room', 'hjy_owner_room.room_id', 'hjy_room.id', 'CASCADE'],
    ]
)

doc.add_heading('3.5 测试数据', level=2)
doc.add_paragraph('系统提供以下初始测试数据：')
add_table(
    ['数据类型', '数量', '说明'],
    [
        ['用户', '2个', 'admin（管理员）、operator（普通用户）'],
        ['角色', '2个', '管理员、普通用户'],
        ['用户角色关联', '2条', 'admin→管理员、operator→普通用户'],
        ['小区', '2个', '阳光花园小区、和谐苑小区'],
        ['楼栋', '3个', '1号楼、2号楼、A座'],
        ['单元', '4个', '1单元、2单元等'],
        ['房屋', '6间', '含不同状态（空置、已入住、装修中）'],
        ['业主', '3人', '张三、李四、王五'],
        ['人房绑定', '4条', '含业主、家属、租客三种关系类型'],
    ]
)

# ========== 4. 非功能需求 ==========
doc.add_heading('4. 非功能需求', level=1)

doc.add_heading('4.1 性能需求', level=2)
doc.add_paragraph('页面响应时间 < 2秒', style='List Bullet')
doc.add_paragraph('分页查询避免全表扫描，使用MyBatis-Plus分页插件', style='List Bullet')
doc.add_paragraph('前端级联选择器数据并行加载（Promise.all），减少等待时间', style='List Bullet')

doc.add_heading('4.2 安全需求', level=2)
doc.add_paragraph('身份认证：路由守卫拦截未登录请求，自动跳转登录页', style='List Bullet')
doc.add_paragraph('权限控制：普通用户无法访问写操作，前端隐藏操作按钮，后端校验请求头', style='List Bullet')
doc.add_paragraph('账号保护：admin账号仅能由admin自身操作，防止越权修改/删除', style='List Bullet')
doc.add_paragraph('数据保护：已入住房屋不可删除；有下级数据的上级记录不可删除', style='List Bullet')
doc.add_paragraph('密码安全：用户信息接口返回时隐藏密码字段', style='List Bullet')

doc.add_heading('4.3 界面需求', level=2)
doc.add_paragraph('使用Element Plus组件库，保持风格统一', style='List Bullet')
doc.add_paragraph('支持主流浏览器（Chrome、Firefox、Edge）', style='List Bullet')
doc.add_paragraph('左侧导航栏 + 顶部标题栏的经典后台布局', style='List Bullet')
doc.add_paragraph('表格支持斑马纹、加载状态', style='List Bullet')
doc.add_paragraph('操作结果通过消息提示反馈用户', style='List Bullet')
doc.add_paragraph('弹窗表单支持表单验证', style='List Bullet')

doc.add_heading('4.4 可靠性需求', level=2)
doc.add_paragraph('数据库操作使用事务（@Transactional）保证原子性', style='List Bullet')
doc.add_paragraph('人房绑定/解绑时同步更新房屋状态，保证数据一致性', style='List Bullet')
doc.add_paragraph('外键约束防止孤立数据产生', style='List Bullet')
doc.add_paragraph('删除操作前进行业务校验，防止误删', style='List Bullet')

# ========== 5. 接口规范 ==========
doc.add_heading('5. 接口规范', level=1)

doc.add_heading('5.1 统一响应格式', level=2)
doc.add_paragraph('所有接口返回统一的JSON格式：')
p = doc.add_paragraph()
run = p.add_run('{\n  "code": 200,\n  "message": "操作成功",\n  "data": { ... }\n}')
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_table(
    ['字段', '类型', '说明'],
    [
        ['code', 'Integer', '状态码，200成功，500失败，403无权限'],
        ['message', 'String', '提示信息'],
        ['data', 'Object', '返回数据'],
    ]
)

doc.add_heading('5.2 分页请求参数', level=2)
add_table(
    ['参数', '类型', '必填', '默认值', '说明'],
    [
        ['pageNum', 'Integer', '否', '1', '当前页码'],
        ['pageSize', 'Integer', '否', '10', '每页条数'],
    ]
)

# ========== 6. 附录 ==========
doc.add_heading('6. 附录', level=1)

doc.add_heading('6.1 术语说明', level=2)
add_table(
    ['术语', '说明'],
    [
        ['CRUD', 'Create/Read/Update/Delete，增删改查'],
        ['RBAC', 'Role-Based Access Control，基于角色的访问控制'],
        ['REST API', 'Representational State Transfer，表述性状态转移风格接口'],
        ['FK', 'Foreign Key，外键约束'],
        ['ORM', 'Object-Relational Mapping，对象关系映射'],
        ['B/S', 'Browser/Server，浏览器/服务器架构'],
    ]
)

doc.add_heading('6.2 参考资料', level=2)
doc.add_paragraph('《数据库系统原理与设计》（第3版），陆鑫等，人民邮电出版社', style='List Bullet')
doc.add_paragraph('Spring Boot官方文档：https://spring.io/projects/spring-boot', style='List Bullet')
doc.add_paragraph('Vue.js官方文档：https://vuejs.org/', style='List Bullet')
doc.add_paragraph('Element Plus组件库文档：https://element-plus.org/', style='List Bullet')
doc.add_paragraph('MyBatis-Plus官方文档：https://baomidou.com/', style='List Bullet')

# ========== 保存 ==========
output_path = r'c:\Users\GWB\Desktop\大三下课程\数据库\合家云\需求分析报告.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
