# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ========== 页面设置 ==========
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ========== 全局样式 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_cell(cell, text, font_size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('北京邮电大学')
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('课程设计报告')
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

# ========== 信息表格 ==========
table = doc.add_table(rows=8, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 第1行：课程名称
set_cell(table.cell(0, 0), '课程名称', bold=True)
set_cell(table.cell(0, 1), '数据库系统原理', bold=False)
table.cell(0, 1).merge(table.cell(0, 3))

# 第2行：题目名称
set_cell(table.cell(1, 0), '题目名称', bold=True)
set_cell(table.cell(1, 1), '合家云社区物业管理平台', bold=False)
table.cell(1, 1).merge(table.cell(1, 3))

# 第3行：学生信息
set_cell(table.cell(2, 0), '姓　　名', bold=True)
set_cell(table.cell(2, 1), '')
set_cell(table.cell(2, 2), '学　　号', bold=True)
set_cell(table.cell(2, 3), '')

# 第4行：学院/专业
set_cell(table.cell(3, 0), '学　　院', bold=True)
set_cell(table.cell(3, 1), '')
set_cell(table.cell(3, 2), '专　　业', bold=True)
set_cell(table.cell(3, 3), '')

# 第5行：班级/指导教师
set_cell(table.cell(4, 0), '班　　级', bold=True)
set_cell(table.cell(4, 1), '')
set_cell(table.cell(4, 2), '指导教师', bold=True)
set_cell(table.cell(4, 3), '')

# 第6行：设计时间
set_cell(table.cell(5, 0), '设计时间', bold=True)
set_cell(table.cell(5, 1), '2026年 3 月 至 2026 年 5 月', bold=False)
table.cell(5, 1).merge(table.cell(5, 3))

# 第7行：成绩
set_cell(table.cell(6, 0), '成　　绩', bold=True)
set_cell(table.cell(6, 1), '')
table.cell(6, 1).merge(table.cell(6, 3))

# 第8行：评语
set_cell(table.cell(7, 0), '指导教师评语', bold=True)
set_cell(table.cell(7, 1), '\n\n\n\n', align=WD_ALIGN_PARAGRAPH.LEFT)
table.cell(7, 1).merge(table.cell(7, 3))

# 设置列宽
for row in table.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(4)
    row.cells[2].width = Cm(3)
    row.cells[3].width = Cm(4)

doc.add_paragraph()

# ========== 课程设计内容摘要 ==========
title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title2.add_run('一、课程设计内容摘要')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

summary_text = (
    '本课题设计并实现了"合家云社区物业管理平台"，采用B/S架构，基于Spring Boot + Vue 3 + MySQL技术栈。'
    '系统实现了社区资产四级层级管理（小区→楼栋→单元→房屋）、业主档案管理、人房绑定管理、'
    'RBAC权限控制（管理员/普通用户）以及数据统计分析等功能。数据库共设计9张表，'
    '包含资产维度（4张）、权限维度（3张）和业主维度（2张），通过外键约束保证数据完整性。'
    '系统支持分页查询、多条件组合查询、级联筛选等数据操作，并通过事务管理保证数据一致性。'
)
p = doc.add_paragraph(summary_text)
p.paragraph_format.first_line_indent = Cm(0.75)

doc.add_paragraph()

# ========== 课程设计成果 ==========
title3 = doc.add_paragraph()
run = title3.add_run('二、课程设计主要成果')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

results = [
    '完成了系统需求分析、总体设计和详细设计，编写了完整的设计文档；',
    '设计了包含9张表的数据库，建立了E-R模型并转化为关系表，设置了外键约束和索引；',
    '实现了后端Spring Boot应用，包含9个Controller、8个Service、8个Mapper，提供RESTful API；',
    '实现了前端Vue 3应用，包含登录页、主页框架、8个管理页面，使用Element Plus组件库；',
    '实现了RBAC权限控制，区分管理员和普通用户角色，支持动态菜单显示和操作权限校验；',
    '实现了数据统计功能，包括概览数据、小区资产统计和房屋状态分布；',
    '完成了系统测试，包括单元测试、接口测试和功能测试，系统运行稳定。',
]
for r in results:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()

# ========== 评分标准 ==========
title4 = doc.add_paragraph()
run = title4.add_run('三、评分标准')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

score_table = doc.add_table(rows=7, cols=3, style='Table Grid')
score_table.alignment = WD_TABLE_ALIGNMENT.CENTER

score_headers = ['评分项目', '分值', '得分']
for i, h in enumerate(score_headers):
    set_cell(score_table.cell(0, i), h, bold=True)

score_items = [
    ['需求分析与设计', '20分', ''],
    ['数据库设计', '20分', ''],
    ['系统功能实现', '30分', ''],
    ['界面设计与交互', '15分', ''],
    ['文档质量', '15分', ''],
    ['合计', '100分', ''],
]
for r_idx, row in enumerate(score_items):
    for c_idx, val in enumerate(row):
        set_cell(score_table.cell(r_idx + 1, c_idx), val,
                 bold=(r_idx == 5))

# 设置列宽
for row in score_table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(3)

# ========== 保存 ==========
output_path = r'c:\Users\GWB\Desktop\大三下课程\数据库\合家云\北京邮电大学课程设计报告表.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
